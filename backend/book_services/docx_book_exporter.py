"""Word (DOCX) book exporter using python-docx.

Implements TOC generation with bookmark hyperlinks and depth-based
indentation, block-type rendering, and the full ``build()`` method stub.
"""

from __future__ import annotations

import io
import re
from lxml import etree

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from book_services.book_exporter import BookExporter, NodeInfo, TocEntry
from book_services.latex_renderer import LatexRenderer


class DocxBookExporter(BookExporter):
    """Export a multi-node lecture book as a Word (.docx) document.

    Task 4.1 scope
    --------------
    - ``__init__``: no special setup required for docx.
    - ``_build_toc``: build ``list[TocEntry]`` from filtered nodes.
    - ``_add_toc_to_doc``: add a TOC section (with bookmark hyperlinks
      and depth-based indentation) to a ``python-docx`` ``Document``.
    - ``build``: stub — raises ``NotImplementedError`` (implemented in 4.3).
    """

    # ------------------------------------------------------------------
    # Initialisation
    # ------------------------------------------------------------------

    def __init__(self) -> None:  # noqa: D401
        """No special setup needed for docx generation."""

    # ------------------------------------------------------------------
    # TOC helpers
    # ------------------------------------------------------------------

    def _build_toc(self, nodes: list[NodeInfo]) -> list[TocEntry]:
        """Build TOC entries from *nodes* (already filtered).

        Anchor names follow the pattern ``node_{depth}_{title}`` as
        specified in the design document.  Page numbers are set to 0
        because Word uses bookmark hyperlinks rather than page numbers.

        Parameters
        ----------
        nodes:
            Ordered list of nodes that all have non-empty ``blocks``.

        Returns
        -------
        list[TocEntry]
            One entry per node, preserving input order.
        """
        entries: list[TocEntry] = []
        for node in nodes:
            anchor = f"node_{node.depth}_{node.text}"
            entries.append(
                TocEntry(
                    title=node.text,
                    depth=node.depth,
                    page=0,  # Word uses bookmarks, not page numbers
                    anchor=anchor,
                )
            )
        return entries

    # ------------------------------------------------------------------
    # XML helpers for bookmarks and hyperlinks
    # ------------------------------------------------------------------

    @staticmethod
    def _make_bookmark_start(bookmark_id: int, name: str):
        """Return a ``<w:bookmarkStart>`` element."""
        el = etree.Element(qn("w:bookmarkStart"))
        el.set(qn("w:id"), str(bookmark_id))
        el.set(qn("w:name"), name)
        return el

    @staticmethod
    def _make_bookmark_end(bookmark_id: int):
        """Return a ``<w:bookmarkEnd>`` element."""
        el = etree.Element(qn("w:bookmarkEnd"))
        el.set(qn("w:id"), str(bookmark_id))
        return el

    @staticmethod
    def _make_hyperlink_run(doc: Document, anchor: str, text: str):
        """Return a ``<w:hyperlink>`` element that jumps to *anchor*.

        Uses an internal bookmark reference (``w:anchor`` attribute)
        rather than an external URL.
        """
        hyperlink = etree.SubElement(
            etree.Element("dummy"),  # temporary parent; will be re-parented
            qn("w:hyperlink"),
        )
        hyperlink = etree.Element(qn("w:hyperlink"))
        hyperlink.set(qn("w:anchor"), anchor)

        run = etree.SubElement(hyperlink, qn("w:r"))

        # Run properties: underline + blue colour to look like a link
        rpr = etree.SubElement(run, qn("w:rPr"))
        color = etree.SubElement(rpr, qn("w:color"))
        color.set(qn("w:val"), "0563C1")
        u = etree.SubElement(rpr, qn("w:u"))
        u.set(qn("w:val"), "single")

        t = etree.SubElement(run, qn("w:t"))
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        return hyperlink

    # ------------------------------------------------------------------
    # TOC section builder
    # ------------------------------------------------------------------

    def _add_toc_to_doc(self, doc: Document, toc_entries: list[TocEntry]) -> None:
        """Add a Table of Contents section to *doc*.

        Each TOC entry is rendered as a paragraph with:
        - Left indentation of ``(depth - 1) * 4`` space-equivalents
          (1 space-equivalent = 0.25 inches / ~360 twips).
        - A hyperlink run that jumps to the corresponding bookmark anchor
          in the document body.

        A "目录" heading paragraph is inserted first, followed by one
        paragraph per entry.

        Parameters
        ----------
        doc:
            The ``python-docx`` ``Document`` to modify in-place.
        toc_entries:
            Ordered list of TOC entries to render.
        """
        # --- TOC heading ---
        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run("目录")
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_para.paragraph_format.space_after = Pt(6)

        # 1 space-equivalent ≈ 0.25 inches in Word indentation terms
        INDENT_UNIT = Inches(0.25)

        for entry in toc_entries:
            para = doc.add_paragraph()

            # Depth-based left indentation: (depth - 1) * 4 space-equivalents
            indent_spaces = (entry.depth - 1) * 4  # number of space-equivalents
            para.paragraph_format.left_indent = indent_spaces * INDENT_UNIT

            # Build hyperlink XML and append to paragraph's XML element
            hyperlink = self._make_hyperlink_run(doc, entry.anchor, entry.title)
            para._p.append(hyperlink)

        # Separator paragraph after TOC
        doc.add_paragraph()

    # ------------------------------------------------------------------
    # LaTeX / text helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _parse_latex_segments(text: str) -> list[dict]:
        """Split *text* into alternating plain-text and LaTeX segments.

        Returns a list of dicts with keys:
        - ``{"type": "text", "content": str}``
        - ``{"type": "latex", "content": str, "display": bool}``

        Display LaTeX (``$$...$$``) is detected before inline (``$...$``)
        to avoid the inner ``$`` being consumed first.
        """
        segments: list[dict] = []
        # 支持 $$...$$, \[...\], \(...\), $...$
        pattern = re.compile(
            r"\$\$(.+?)\$\$"
            r"|\\\[(.+?)\\\]"
            r"|\\\((.+?)\\\)"
            r"|\$(.+?)\$",
            re.DOTALL,
        )
        last_end = 0
        for m in pattern.finditer(text):
            if m.start() > last_end:
                segments.append({"type": "text", "content": text[last_end : m.start()]})
            if m.group(1) is not None:
                segments.append({"type": "latex", "content": m.group(1), "display": True})
            elif m.group(2) is not None:
                segments.append({"type": "latex", "content": m.group(2), "display": True})
            elif m.group(3) is not None:
                segments.append({"type": "latex", "content": m.group(3), "display": False})
            else:
                segments.append({"type": "latex", "content": m.group(4), "display": False})
            last_end = m.end()
        if last_end < len(text):
            segments.append({"type": "text", "content": text[last_end:]})
        return segments

    # ------------------------------------------------------------------
    # Block rendering
    # ------------------------------------------------------------------

    @staticmethod
    def _set_run_font(run, font_name: str = "宋体") -> None:
        """Set both ASCII and East-Asian font on *run* to *font_name*."""
        run.font.name = font_name
        # python-docx requires explicit East-Asian font via XML
        rpr = run._r.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rpr, qn("w:rFonts"))
        rFonts.set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _add_paragraph_shading(para, fill: str = "D9D9D9") -> None:
        """Apply solid background shading to *para* via XML."""
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn("w:shd"))
        if shd is None:
            shd = etree.SubElement(pPr, qn("w:shd"))
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

    @staticmethod
    def _add_paragraph_left_border(para, color: str = "808080", sz: str = "18") -> None:
        """Add a left border to *para* to simulate a block-quote style."""
        pPr = para._p.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        left = pBdr.find(qn("w:left"))
        if left is None:
            left = etree.SubElement(pBdr, qn("w:left"))
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), sz)       # border width in eighths of a point
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), color)

    def _add_text_with_latex(
        self,
        para,
        text: str,
        font_name: str,
        font_size_pt: float,
        latex_renderer: LatexRenderer,
    ) -> None:
        """Append *text* (which may contain LaTeX) to *para*.

        Inline/display LaTeX segments are rendered to PNG and embedded as
        inline pictures.  Falls back to raw text + label on render failure.
        """
        segments = self._parse_latex_segments(text)
        for seg in segments:
            if seg["type"] == "text":
                run = para.add_run(seg["content"])
                run.font.size = Pt(font_size_pt)
                self._set_run_font(run, font_name)
            else:
                png_bytes = latex_renderer.render(seg["content"], display=seg["display"])
                if png_bytes is None:
                    fallback = f"[公式渲染失败，原始代码如下] ${seg['content']}$"
                    run = para.add_run(fallback)
                    run.font.size = Pt(font_size_pt)
                    self._set_run_font(run, font_name)
                else:
                    # Embed PNG as inline picture; scale to ~line height
                    img_buf = io.BytesIO(png_bytes)
                    height_inches = font_size_pt / 72.0 * 1.4  # approx line height
                    run = para.add_run()
                    run.add_picture(img_buf, height=Inches(height_inches))

    def _render_block(
        self,
        doc: Document,
        block: dict,
        latex_renderer: LatexRenderer,
    ) -> None:
        """Render a single *block* dict into *doc*.

        Supported block types: heading, paragraph, code, list, quote.
        Unknown types fall back to Normal paragraph style.

        Requirements: 5.2, 5.3, 5.4, 5.5, 6.1, 6.2, 6.3
        """
        block_type = block.get("type", "paragraph")
        content = block.get("content", "") or block.get("text", "") or ""

        if block_type == "heading":
            level = block.get("level", 1)
            level = max(1, min(3, level))  # clamp to 1–3
            style_name = f"Heading {level}"
            para = doc.add_paragraph(style=style_name)
            run = para.add_run(content)
            self._set_run_font(run, "宋体")

        elif block_type == "paragraph":
            para = doc.add_paragraph(style="Normal")
            self._add_text_with_latex(para, content, "宋体", 11, latex_renderer)

        elif block_type == "code":
            para = doc.add_paragraph(style="Normal")
            self._add_paragraph_shading(para, fill="D9D9D9")
            run = para.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            # Ensure East-Asian font is also set to Courier New for consistency
            rpr = run._r.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = etree.SubElement(rpr, qn("w:rFonts"))
            rFonts.set(qn("w:eastAsia"), "Courier New")

        elif block_type == "list":
            para = doc.add_paragraph(style="List Bullet")
            self._add_text_with_latex(para, content, "宋体", 11, latex_renderer)

        elif block_type == "quote":
            para = doc.add_paragraph(style="Normal")
            para.paragraph_format.left_indent = Inches(0.25)
            self._add_paragraph_left_border(para)
            self._add_text_with_latex(para, content, "宋体", 11, latex_renderer)

        else:
            # Unknown type — render as Normal paragraph
            para = doc.add_paragraph(style="Normal")
            self._add_text_with_latex(para, content, "宋体", 11, latex_renderer)

    # ------------------------------------------------------------------
    # Main build method
    # ------------------------------------------------------------------

    def build(
        self,
        session_title: str,
        nodes: list[NodeInfo],
        include_toc: bool = True,
    ) -> bytes:
        """Build a Word (.docx) document from *nodes* and return raw bytes.

        Steps:
        1. Instantiate a fresh ``LatexRenderer`` (per-call cache scope).
        2. Filter nodes to those with non-empty ``blocks``.
        3. Optionally prepend a TOC section.
        4. Iterate filtered nodes; render each block; insert a section break
           between nodes.
        5. Serialise the document to ``io.BytesIO`` and return the bytes.

        Requirements: 5.1
        """
        latex_renderer = LatexRenderer()
        filtered = self._filter_nodes(nodes)

        doc = Document()

        # Optional TOC
        if include_toc:
            toc_entries = self._build_toc(filtered)
            self._add_toc_to_doc(doc, toc_entries)

        # Render each node's blocks
        for node_index, node in enumerate(filtered):
            # Add a bookmark anchor paragraph for TOC hyperlinks
            anchor_name = f"node_{node.depth}_{node.text}"
            anchor_para = doc.add_paragraph()
            bookmark_id = node_index
            anchor_para._p.insert(
                0, self._make_bookmark_start(bookmark_id, anchor_name)
            )
            anchor_para._p.append(self._make_bookmark_end(bookmark_id))

            # Render all blocks for this node
            for block in node.blocks:
                self._render_block(doc, block, latex_renderer)

            # Insert a section break between nodes (not after the last one)
            if node_index < len(filtered) - 1:
                last_para = doc.add_paragraph()
                last_para_pPr = last_para._p.get_or_add_pPr()
                sectPr = etree.SubElement(last_para_pPr, qn("w:sectPr"))
                sectType = etree.SubElement(sectPr, qn("w:type"))
                sectType.set(qn("w:val"), "nextPage")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
