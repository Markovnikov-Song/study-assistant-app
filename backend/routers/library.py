"""
library.py — 学校/图书馆路由模块
挂载在 /api/library
"""
from __future__ import annotations

import re
from typing import Any, Literal, Optional

from fastapi import APIRouter, Depends, HTTPException
from pydantic import BaseModel, field_validator
from sqlalchemy import text

from database import (
    ConversationHistory,
    ConversationSession,
    MindmapNodeState,
    NodeLecture,
    Subject,
    get_session as db_session,
)
from deps import get_current_user

router = APIRouter()


# ---------------------------------------------------------------------------
# Pydantic schemas
# ---------------------------------------------------------------------------


class SubjectProgressOut(BaseModel):
    id: int
    name: str
    category: Optional[str]
    is_pinned: int
    session_count: int
    total_nodes: int
    lit_nodes: int
    last_visited_at: Optional[str]


class SessionSummaryOut(BaseModel):
    id: int
    title: Optional[str]
    created_at: str
    total_nodes: int
    lit_nodes: int
    is_pinned: bool
    sort_order: int


class SessionMetaIn(BaseModel):
    is_pinned: Optional[bool] = None
    sort_order: Optional[int] = None


class SessionMetaOut(BaseModel):
    id: int
    is_pinned: bool
    sort_order: int


class TitleIn(BaseModel):
    title: str

    @field_validator("title")
    @classmethod
    def validate_title(cls, v: str) -> str:
        v = v.strip()
        if not v:
            raise ValueError("标题不能为空")
        if len(v) > 64:
            raise ValueError("标题不能超过 64 个字符")
        return v


class NodeStateItem(BaseModel):
    node_id: str
    is_lit: bool


class NodeStatesIn(BaseModel):
    states: list[NodeStateItem]


class LectureContentIn(BaseModel):
    content: dict[str, Any]


class ContentIn(BaseModel):
    content: str


class ExportBookIn(BaseModel):
    node_ids: list[str]
    format: Literal["pdf", "docx"]
    include_toc: bool = True

    @field_validator("node_ids")
    @classmethod
    def validate_node_ids(cls, v: list[str]) -> list[str]:
        if not v:
            raise ValueError("node_ids 不能为空")
        return v

    @field_validator("format")
    @classmethod
    def validate_format(cls, v: str) -> str:
        if v not in ("pdf", "docx"):
            raise ValueError("不支持的导出格式")
        return v


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _parse_node_count(markdown: str) -> int:
    """统计 Markdown 中 # 开头的节点行数。"""
    return sum(1 for line in markdown.splitlines() if re.match(r"^#{1,4}\s+\S", line))


def _get_mindmap_content(db, session_id: int) -> Optional[str]:
    """从 conversation_history 取该会话最新 assistant 记录的 content。"""
    record = (
        db.query(ConversationHistory)
        .filter_by(session_id=session_id, role="assistant")
        .order_by(ConversationHistory.created_at.desc())
        .first()
    )
    return record.content if record else None


def _assert_session_owner(db, session_id: int, user_id: int) -> ConversationSession:
    sess = db.query(ConversationSession).filter_by(id=session_id, user_id=user_id).first()
    if not sess:
        raise HTTPException(404, "大纲不存在")
    return sess


# ---------------------------------------------------------------------------
# 学科列表
# ---------------------------------------------------------------------------


@router.get("/subjects", response_model=list[SubjectProgressOut])
def get_subjects(user=Depends(get_current_user)):
    with db_session() as db:
        subjects = (
            db.query(Subject)
            .filter_by(user_id=user["id"], is_archived=0)
            .order_by(Subject.is_pinned.desc(), Subject.created_at.desc())
            .all()
        )
        if not subjects:
            return []

        subject_ids = [s.id for s in subjects]

        # 批量查询所有相关 mindmap sessions
        all_sessions = (
            db.query(ConversationSession)
            .filter(
                ConversationSession.user_id == user["id"],
                ConversationSession.subject_id.in_(subject_ids),
                ConversationSession.session_type == "mindmap",
            )
            .order_by(
                ConversationSession.is_pinned.desc(),
                ConversationSession.created_at.desc(),
            )
            .all()
        )

        # 按 subject_id 分组，取每个学科的 active session（置顶优先，否则最新）
        from collections import defaultdict
        sessions_by_subject: dict[int, list] = defaultdict(list)
        for s in all_sessions:
            sessions_by_subject[s.subject_id].append(s)

        active_session_ids = [
            sessions_by_subject[sid][0].id
            for sid in subject_ids
            if sessions_by_subject[sid]
        ]

        # 批量查询 mindmap 内容（最新 assistant 消息）
        mindmap_contents: dict[int, str] = {}
        if active_session_ids:
            # 用子查询取每个 session 最新的 assistant 消息
            from sqlalchemy import func
            latest_ids = (
                db.query(func.max(ConversationHistory.id))
                .filter(
                    ConversationHistory.session_id.in_(active_session_ids),
                    ConversationHistory.role == "assistant",
                )
                .group_by(ConversationHistory.session_id)
                .subquery()
            )
            records = (
                db.query(ConversationHistory)
                .filter(ConversationHistory.id.in_(latest_ids))
                .all()
            )
            mindmap_contents = {r.session_id: r.content for r in records}

        # 批量查询点亮节点数
        lit_counts: dict[int, int] = {}
        if active_session_ids:
            from sqlalchemy import func
            rows = (
                db.query(
                    MindmapNodeState.session_id,
                    func.count(MindmapNodeState.id).label("cnt"),
                )
                .filter(
                    MindmapNodeState.user_id == user["id"],
                    MindmapNodeState.session_id.in_(active_session_ids),
                    MindmapNodeState.is_lit == 1,
                )
                .group_by(MindmapNodeState.session_id)
                .all()
            )
            lit_counts = {r.session_id: r.cnt for r in rows}

        result = []
        for subj in subjects:
            sess_list = sessions_by_subject[subj.id]
            active = sess_list[0] if sess_list else None
            total_nodes = 0
            lit_nodes = 0
            last_visited_at = None
            if active:
                content = mindmap_contents.get(active.id)
                if content:
                    total_nodes = _parse_node_count(content)
                lit_nodes = lit_counts.get(active.id, 0)
                last_visited_at = active.created_at
            result.append(SubjectProgressOut(
                id=subj.id,
                name=subj.name,
                category=subj.category,
                is_pinned=subj.is_pinned,
                session_count=len(sess_list),
                total_nodes=total_nodes,
                lit_nodes=lit_nodes,
                last_visited_at=last_visited_at.isoformat() if last_visited_at else None,
            ))
        return result


# ---------------------------------------------------------------------------
# 学科下的大纲列表
# ---------------------------------------------------------------------------


@router.get("/subjects/{subject_id}/sessions", response_model=list[SessionSummaryOut])
def get_sessions(subject_id: int, user=Depends(get_current_user)):
    with db_session() as db:
        sessions = (
            db.query(ConversationSession)
            .filter_by(user_id=user["id"], subject_id=subject_id, session_type="mindmap")
            .order_by(
                ConversationSession.is_pinned.desc(),
                ConversationSession.sort_order.asc(),
                ConversationSession.created_at.desc(),
            )
            .all()
        )
        if not sessions:
            return []

        session_ids = [s.id for s in sessions]

        # 批量查 mindmap 内容
        from sqlalchemy import func
        latest_ids = (
            db.query(func.max(ConversationHistory.id))
            .filter(
                ConversationHistory.session_id.in_(session_ids),
                ConversationHistory.role == "assistant",
            )
            .group_by(ConversationHistory.session_id)
            .subquery()
        )
        records = (
            db.query(ConversationHistory)
            .filter(ConversationHistory.id.in_(latest_ids))
            .all()
        )
        content_map = {r.session_id: r.content for r in records}

        # 批量查点亮数
        lit_rows = (
            db.query(
                MindmapNodeState.session_id,
                func.count(MindmapNodeState.id).label("cnt"),
            )
            .filter(
                MindmapNodeState.user_id == user["id"],
                MindmapNodeState.session_id.in_(session_ids),
                MindmapNodeState.is_lit == 1,
            )
            .group_by(MindmapNodeState.session_id)
            .all()
        )
        lit_map = {r.session_id: r.cnt for r in lit_rows}

        return [
            SessionSummaryOut(
                id=sess.id,
                title=sess.title,
                created_at=sess.created_at.isoformat(),
                total_nodes=_parse_node_count(content_map.get(sess.id, "")),
                lit_nodes=lit_map.get(sess.id, 0),
                is_pinned=bool(getattr(sess, "is_pinned", 0)),
                sort_order=getattr(sess, "sort_order", 0),
            )
            for sess in sessions
        ]


# ---------------------------------------------------------------------------
# 更新大纲元数据（置顶 / 排序）
# ---------------------------------------------------------------------------


@router.patch("/sessions/{session_id}/meta", response_model=SessionMetaOut)
def update_session_meta(session_id: int, body: SessionMetaIn, user=Depends(get_current_user)):
    with db_session() as db:
        sess = _assert_session_owner(db, session_id, user["id"])
        if body.is_pinned is not None:
            sess.is_pinned = 1 if body.is_pinned else 0
        if body.sort_order is not None:
            sess.sort_order = body.sort_order
        db.flush()
        return SessionMetaOut(
            id=sess.id,
            is_pinned=bool(sess.is_pinned),
            sort_order=sess.sort_order,
        )


# ---------------------------------------------------------------------------
# 重命名大纲
# ---------------------------------------------------------------------------


@router.patch("/sessions/{session_id}/title")
def rename_session(session_id: int, body: TitleIn, user=Depends(get_current_user)):
    with db_session() as db:
        sess = _assert_session_owner(db, session_id, user["id"])
        sess.title = body.title
    return {"ok": True}


# ---------------------------------------------------------------------------
# 删除大纲（级联删除讲义和点亮状态）
# ---------------------------------------------------------------------------


@router.delete("/sessions/{session_id}")
def delete_session(session_id: int, user=Depends(get_current_user)):
    with db_session() as db:
        sess = _assert_session_owner(db, session_id, user["id"])
        # cascade via FK ON DELETE CASCADE handles node_states and lectures
        db.delete(sess)
    return {"ok": True}


# ---------------------------------------------------------------------------
# 节点树（解析 Markdown 返回节点树 JSON）
# ---------------------------------------------------------------------------


def _build_node_tree(markdown: str) -> list[dict]:
    """将 Markdown 标题解析为节点树列表（扁平，含 parent_id）。"""
    nodes: list[dict] = []
    ancestor_stack: list[dict] = []  # stack of (depth, node)
    sibling_counter: dict[str, int] = {}  # key -> count

    for line in markdown.splitlines():
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if not m:
            continue
        depth = len(m.group(1))
        text = m.group(2).strip()
        if not text:
            continue

        # pop stack to find parent
        while ancestor_stack and ancestor_stack[-1]["depth"] >= depth:
            ancestor_stack.pop()

        parent = ancestor_stack[-1] if ancestor_stack else None
        parent_id = parent["node_id"] if parent else None

        # build ancestor path for node_id
        if parent:
            ancestor_path = parent["node_id"].split("_", 1)[1] if "_" in parent["node_id"] else parent["node_id"]
            base_key = f"L{depth}_{ancestor_path}_{text}"
        else:
            base_key = f"L{depth}_{text}"

        count = sibling_counter.get(base_key, 0) + 1
        sibling_counter[base_key] = count
        node_id = base_key if count == 1 else f"{base_key}_{count}"

        node = {
            "node_id": node_id,
            "text": text,
            "depth": depth,
            "parent_id": parent_id,
            "is_user_created": False,
            "children": [],
        }
        nodes.append(node)
        ancestor_stack.append({"depth": depth, "node_id": node_id})

    return nodes


@router.get("/sessions/{session_id}/nodes")
def get_nodes(session_id: int, user=Depends(get_current_user)):
    with db_session() as db:
        _assert_session_owner(db, session_id, user["id"])
        content = _get_mindmap_content(db, session_id)
        if not content:
            return {"nodes": []}
        return {"nodes": _build_node_tree(content)}


# ---------------------------------------------------------------------------
# 更新大纲 Markdown 内容
# ---------------------------------------------------------------------------


@router.patch("/sessions/{session_id}/content")
def update_content(session_id: int, body: ContentIn, user=Depends(get_current_user)):
    with db_session() as db:
        _assert_session_owner(db, session_id, user["id"])
        record = (
            db.query(ConversationHistory)
            .filter_by(session_id=session_id, role="assistant")
            .order_by(ConversationHistory.created_at.desc())
            .first()
        )
        if record:
            record.content = body.content
        else:
            db.add(ConversationHistory(session_id=session_id, role="assistant", content=body.content))
    return {"ok": True}


# ---------------------------------------------------------------------------
# 节点点亮状态
# ---------------------------------------------------------------------------


@router.get("/sessions/{session_id}/node-states")
def get_node_states(session_id: int, user=Depends(get_current_user)):
    with db_session() as db:
        _assert_session_owner(db, session_id, user["id"])
        rows = (
            db.query(MindmapNodeState)
            .filter_by(user_id=user["id"], session_id=session_id)
            .all()
        )
        return {row.node_id: bool(row.is_lit) for row in rows}


@router.post("/sessions/{session_id}/node-states")
def upsert_node_states(session_id: int, body: NodeStatesIn, user=Depends(get_current_user)):
    with db_session() as db:
        _assert_session_owner(db, session_id, user["id"])
        for item in body.states:
            existing = (
                db.query(MindmapNodeState)
                .filter_by(user_id=user["id"], session_id=session_id, node_id=item.node_id)
                .first()
            )
            if existing:
                existing.is_lit = 1 if item.is_lit else 0
            else:
                db.add(
                    MindmapNodeState(
                        user_id=user["id"],
                        session_id=session_id,
                        node_id=item.node_id,
                        is_lit=1 if item.is_lit else 0,
                    )
                )

        # 静默同步：节点点亮时自动更新对应 plan_item 状态
        lit_node_ids = [item.node_id for item in body.states if item.is_lit]
        if lit_node_ids:
            try:
                from services.study_planner_service import StudyPlannerService
                svc = StudyPlannerService()
                for node_id in lit_node_ids:
                    svc.sync_node_completion(
                        user_id=int(user["id"]),
                        node_id=node_id,
                        db=db,
                    )
            except Exception as _sync_err:
                import logging
                logging.getLogger(__name__).warning(
                    "study_planner sync_node_completion 失败（非致命）：%s", _sync_err
                )

    return {"ok": True}


# ---------------------------------------------------------------------------
# Session 三层学习进度
# ---------------------------------------------------------------------------
#
# 学习链条：
#   阅读层  (read)     — 节点点亮率（is_lit），权重 0.3
#   练习层  (practice) — 该 session 节点关联的 SM-2 卡片完成率，权重 0.5
#   掌握层  (mastery)  — SM-2 mastery_score 均值（0-5 → 0.0-1.0），权重 0.2
#
# 进度描述对象是单个 session（讲义/导图），不是学科。
# ---------------------------------------------------------------------------


class SessionProgressOut(BaseModel):
    session_id: int
    total_nodes: int
    lit_nodes: int          # 已点亮节点数
    lecture_nodes: int      # 已生成讲义的节点数
    # 三层进度 0.0~1.0
    read_progress: float    # 阅读层：lit_nodes / total_nodes
    practice_progress: float  # 练习层：已复习卡片 / 总卡片
    mastery_progress: float   # 掌握层：avg mastery_score / 5
    overall_progress: float   # 综合：read×0.3 + practice×0.5 + mastery×0.2
    # 错题统计
    mistake_count: int
    reviewed_mistake_count: int


@router.get("/sessions/{session_id}/progress", response_model=SessionProgressOut)
def get_session_progress(session_id: int, user=Depends(get_current_user)):
    """
    返回单个 session 的三层学习进度。
    - 阅读层：节点点亮率
    - 练习层：该 session 节点关联的 SM-2 复习卡片完成情况
    - 掌握层：SM-2 mastery_score 均值
    """
    from database import NodeLecture, ReviewCard, Note as NoteModel

    with db_session() as db:
        sess = _assert_session_owner(db, session_id, user["id"])

        # ── 阅读层 ──────────────────────────────────────────────────────────
        node_states = (
            db.query(MindmapNodeState)
            .filter_by(user_id=user["id"], session_id=session_id)
            .all()
        )
        lit_node_ids = {s.node_id for s in node_states if s.is_lit}

        # 从 session content 字段解析总节点数
        content_row = db.execute(
            text("SELECT content FROM conversation_sessions WHERE id = :sid"),
            {"sid": session_id},
        ).fetchone()
        total_nodes = _parse_node_count(content_row[0] if content_row and content_row[0] else "")
        # fallback: 用已知节点数估算
        if total_nodes == 0:
            total_nodes = max(len(node_states), 1)

        lit_nodes = len(lit_node_ids)
        read_progress = lit_nodes / total_nodes if total_nodes > 0 else 0.0

        # ── 讲义层（已生成讲义的节点数）────────────────────────────────────
        lecture_nodes = (
            db.query(NodeLecture)
            .filter_by(user_id=user["id"], session_id=session_id)
            .count()
        )

        # ── 练习层 + 掌握层：通过 node_id 关联 ReviewCard ──────────────────
        # ReviewCard.node_id 与 MindmapNodeState.node_id 一致
        if lit_node_ids:
            review_cards = (
                db.query(ReviewCard)
                .filter(
                    ReviewCard.user_id == user["id"],
                    ReviewCard.node_id.in_(lit_node_ids),
                )
                .all()
            )
        else:
            review_cards = []

        total_cards = len(review_cards)
        if total_cards > 0:
            # 已复习 = total_reviews > 0
            reviewed_cards = sum(1 for c in review_cards if c.total_reviews > 0)
            # 练习进度：已复习卡片占比，加 0.3 基础分（有卡片就算开始练习）
            practice_progress = min(
                (reviewed_cards + total_cards * 0.3) / (total_cards * 1.3), 1.0
            )
            avg_mastery = sum(c.mastery_score for c in review_cards) / total_cards
            mastery_progress = avg_mastery / 5.0
        else:
            practice_progress = 0.0
            mastery_progress = 0.0

        overall_progress = (
            read_progress * 0.3
            + practice_progress * 0.5
            + mastery_progress * 0.2
        )

        # ── 错题统计 ────────────────────────────────────────────────────────
        # 通过 node_id 关联该 session 的错题
        if lit_node_ids:
            mistake_notes = (
                db.query(NoteModel)
                .filter(
                    NoteModel.user_id == user["id"],
                    NoteModel.note_type == "mistake",
                    NoteModel.node_id.in_(lit_node_ids),
                )
                .all()
            )
        else:
            mistake_notes = []

        mistake_count = len(mistake_notes)
        reviewed_mistake_count = sum(
            1 for n in mistake_notes if n.mistake_status == "reviewed"
        )

        return SessionProgressOut(
            session_id=session_id,
            total_nodes=total_nodes,
            lit_nodes=lit_nodes,
            lecture_nodes=lecture_nodes,
            read_progress=round(read_progress, 3),
            practice_progress=round(practice_progress, 3),
            mastery_progress=round(mastery_progress, 3),
            overall_progress=round(overall_progress, 3),
            mistake_count=mistake_count,
            reviewed_mistake_count=reviewed_mistake_count,
        )
# ---------------------------------------------------------------------------


@router.get("/lectures/{session_id}")
def get_lecture(session_id: int, node_id: str, user=Depends(get_current_user)):
    from urllib.parse import unquote_plus
    # 前端可能把空格编码为 +，FastAPI 默认不解码 +，需要手动处理
    node_id = unquote_plus(node_id)
    with db_session() as db:
        lecture = (
            db.query(NodeLecture)
            .filter_by(user_id=user["id"], session_id=session_id, node_id=node_id)
            .first()
        )
        if not lecture:
            raise HTTPException(404, "讲义不存在")
        return {
            "id": lecture.id,
            "node_id": lecture.node_id,
            "content": lecture.content,
            "resource_scope": lecture.resource_scope,
            "created_at": lecture.created_at.isoformat(),
            "updated_at": lecture.updated_at.isoformat(),
        }


class LectureCreateIn(BaseModel):
    session_id: int
    node_id: str
    content: Optional[dict[str, Any]] = None  # ignored; LLM generates content
    resource_scope: Optional[dict[str, Any]] = None


@router.post("/lectures")
def create_lecture(body: LectureCreateIn, user=Depends(get_current_user)):
    """生成节点讲义（一次性，兼容旧接口）。超时返回 504。"""
    import concurrent.futures
    from services.lecture_generator_service import LectureGeneratorService

    with db_session() as db:
        session = _assert_session_owner(db, body.session_id, user["id"])
        subject_id: int = session.subject_id or 0

    node_title = _parse_text_from_node_id(body.node_id)
    svc = LectureGeneratorService()

    def _generate():
        return svc.generate(
            session_id=body.session_id,
            node_id=body.node_id,
            node_title=node_title,
            user_id=user["id"],
            subject_id=subject_id,
            resource_scope=body.resource_scope,
        )

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(_generate)
            from backend_config import get_config
            result = future.result(timeout=get_config().LIBRARY_LECTURE_GENERATE_TIMEOUT_SECONDS)
        return {"id": result["id"], "ok": True}
    except concurrent.futures.TimeoutError:
        raise HTTPException(504, "讲义生成超时，请重试")
    except RuntimeError as e:
        raise HTTPException(502, f"AI 服务暂时不可用：{e}")
    except Exception as e:
        raise HTTPException(500, f"讲义生成失败：{e}")


@router.post("/lectures/stream")
def create_lecture_stream(body: LectureCreateIn, user=Depends(get_current_user)):
    """流式生成讲义，返回 SSE。每个 token 以 data: <token>\\n\\n 发送，结束发 data: [DONE]\\n\\n。"""
    from fastapi.responses import StreamingResponse
    from services.lecture_generator_service import LectureGeneratorService

    with db_session() as db:
        session = _assert_session_owner(db, body.session_id, user["id"])
        subject_id: int = session.subject_id or 0

    node_title = _parse_text_from_node_id(body.node_id)
    svc = LectureGeneratorService()
    user_id: int = user["id"]

    def event_generator():
        accumulated = []
        try:
            for token in svc.generate_stream(
                node_id=body.node_id,
                node_title=node_title,
                subject_id=subject_id,
                session_id=body.session_id,
                resource_scope=body.resource_scope,
            ):
                accumulated.append(token)
                # 换行符转义，避免破坏 SSE 帧格式
                safe_token = token.replace('\n', '\\n')
                yield f"data: {safe_token}\n\n"
            # 流结束后保存完整内容到数据库
            full_markdown = "".join(accumulated)
            if full_markdown.strip():
                svc.save_stream_result(
                    full_markdown=full_markdown,
                    node_id=body.node_id,
                    session_id=body.session_id,
                    user_id=user_id,
                    resource_scope=body.resource_scope,
                )
        except Exception as e:
            import traceback; traceback.print_exc()
            yield f"data: [ERROR]{e}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.patch("/lectures/{lecture_id}")
def patch_lecture(lecture_id: int, body: LectureContentIn, user=Depends(get_current_user)):
    with db_session() as db:
        lecture = db.query(NodeLecture).filter_by(id=lecture_id, user_id=user["id"]).first()
        if not lecture:
            raise HTTPException(404, "讲义不存在")
        lecture.content = body.content
    return {"ok": True}


@router.delete("/lectures/{session_id}")
def delete_lecture(session_id: int, node_id: str, user=Depends(get_current_user)):
    from urllib.parse import unquote_plus
    node_id = unquote_plus(node_id)
    with db_session() as db:
        lecture = (
            db.query(NodeLecture)
            .filter_by(user_id=user["id"], session_id=session_id, node_id=node_id)
            .first()
        )
        if not lecture:
            raise HTTPException(404, "讲义不存在")
        db.delete(lecture)
    return {"ok": True}


@router.post("/lectures/{lecture_id}/export")
def export_lecture(lecture_id: int, format: str = "docx", user=Depends(get_current_user)):
    """导出讲义为 PDF 或 Word 文件。"""
    from backend_config import get_config
    if format not in ("pdf", "docx"):
        raise HTTPException(400, "不支持的格式，请使用 pdf 或 docx")

    with db_session() as db:
        lecture = db.query(NodeLecture).filter_by(id=lecture_id, user_id=user["id"]).first()
        if not lecture:
            raise HTTPException(404, "讲义不存在")
        content = lecture.content
        node_id = lecture.node_id

    blocks = (content or {}).get("blocks", [])
    if not blocks:
        raise HTTPException(422, "讲义内容为空，无法导出")
    # 取节点名作为标题（从 node_id 解析最后一段）
    title = _parse_text_from_node_id(node_id) if node_id else "讲义"

    from fastapi.responses import StreamingResponse
    import io

    # ── PDF ──────────────────────────────────────────────────────────────────
    if format == "pdf":
        try:
            from book_services.pdf_book_exporter import PdfBookExporter
            from book_services.book_exporter import NodeInfo
        except ImportError as e:
            raise HTTPException(500, f"PDF 导出依赖缺失：{e}")

        import concurrent.futures

        def _build_pdf():
            exporter = PdfBookExporter()
            node_info = NodeInfo(
                node_id=node_id,
                text=title,
                depth=1,
                blocks=blocks,
            )
            return exporter.build(
                session_title=title,
                nodes=[node_info],
                include_toc=False,
            )

        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                future = pool.submit(_build_pdf)
                pdf_bytes = future.result(timeout=get_config().LIBRARY_PDF_EXPORT_TIMEOUT_SECONDS)
        except concurrent.futures.TimeoutError:
            raise HTTPException(504, "PDF 生成超时，请稍后重试")
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(500, f"PDF 生成失败：{type(e).__name__}: {e}")

        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename=lecture_{lecture_id}.pdf"},
        )

    # ── DOCX ─────────────────────────────────────────────────────────────────
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise HTTPException(500, "python-docx 未安装")

    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for block in blocks:
        btype = block.get("type", "paragraph")
        text = block.get("text", "")
        if btype == "heading":
            level = block.get("level", 2)
            doc.add_heading(text, level=level)
        elif btype == "code":
            p = doc.add_paragraph(text)
            p.style = "No Spacing"
        elif btype == "list":
            doc.add_paragraph(text, style="List Bullet")
        elif btype == "quote":
            doc.add_paragraph(f"> {text}")
        else:
            doc.add_paragraph(text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=lecture_{lecture_id}.docx"},
    )


# ---------------------------------------------------------------------------
# 导出为书
# ---------------------------------------------------------------------------


def _parse_depth_from_node_id(node_id: str) -> int:
    """Extract depth from node_id if it follows the L{depth}_... pattern."""
    import re as _re
    m = _re.match(r"^L(\d+)_", node_id)
    if m:
        return int(m.group(1))
    return 1


def _parse_text_from_node_id(node_id: str) -> str:
    """Extract display title from node_id.

    node_id format: L{depth}_{ancestor_path}_{title}
    e.g. "L1_基本概念" → "基本概念"
         "L2_第1章_绪论_基本概念" → "基本概念"
    We take the last underscore-separated segment as the title.
    """
    import re as _re
    # Strip leading L{n}_ prefix
    text = _re.sub(r"^L\d+_", "", node_id)
    # The last segment after splitting by _ is the node's own title
    # But node titles themselves may contain underscores, so we can't just split.
    # Best effort: return the full text after stripping the depth prefix.
    return text or node_id


@router.post("/sessions/{session_id}/export-book")
def export_book(session_id: int, body: ExportBookIn, user=Depends(get_current_user)):
    """将多个节点的讲义合并导出为 PDF 或 Word 书籍。

    Requirements: 4.1, 4.6, 5.1, 5.6, 7.1, 7.5, 7.6
    """
    import io as _io
    import sys as _sys
    import os as _os
    from fastapi.responses import StreamingResponse
    from book_services.book_exporter import NodeInfo
    from book_services.pdf_book_exporter import PdfBookExporter
    from book_services.docx_book_exporter import DocxBookExporter

    with db_session() as db:
        # 1. Verify session ownership
        sess = _assert_session_owner(db, session_id, user["id"])
        session_title: str = sess.title or f"session_{session_id}"

        # 2. 从 mindmap 内容中建立 node_id → 真实标题 的映射
        mindmap_content = _get_mindmap_content(db, session_id)
        node_text_map: dict[str, str] = {}
        if mindmap_content:
            for node_dict in _build_node_tree(mindmap_content):
                node_text_map[node_dict["node_id"]] = node_dict["text"]

        # 3. Batch-query node_lectures for the requested node_ids
        lectures = (
            db.query(NodeLecture)
            .filter(
                NodeLecture.user_id == user["id"],
                NodeLecture.session_id == session_id,
                NodeLecture.node_id.in_(body.node_ids),
            )
            .all()
        )

        # Build a lookup map for O(1) access
        lecture_map: dict[str, NodeLecture] = {lec.node_id: lec for lec in lectures}

        # 4. Preserve input order; build NodeInfo objects; filter nodes without content
        nodes: list[NodeInfo] = []
        for nid in body.node_ids:
            lec = lecture_map.get(nid)
            if lec is None:
                continue
            blocks = (lec.content or {}).get("blocks", []) if lec.content else []
            if not blocks:
                continue
            # 优先用 mindmap 里的真实标题，fallback 到解析 node_id
            display_text = node_text_map.get(nid) or _parse_text_from_node_id(nid)
            nodes.append(
                NodeInfo(
                    node_id=nid,
                    text=display_text,
                    depth=_parse_depth_from_node_id(nid),
                    blocks=blocks,
                )
            )

    # 4. All nodes filtered out → 422
    if not nodes:
        raise HTTPException(422, "所选节点均无讲义内容")

    # 5. Instantiate exporter and build document
    ext = body.format  # "pdf" or "docx"
    try:
        if body.format == "pdf":
            exporter = PdfBookExporter()
            media_type = "application/pdf"
        else:
            exporter = DocxBookExporter()
            media_type = (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            )
    except RuntimeError as e:
        raise HTTPException(500, str(e))

    try:
        file_bytes = exporter.build(
            session_title=session_title,
            nodes=nodes,
            include_toc=body.include_toc,
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(500, f"导出失败：{type(e).__name__}: {e}")

    # 6. Return streaming response with Content-Disposition
    filename = f"book_{session_id}.{ext}"
    return StreamingResponse(
        _io.BytesIO(file_bytes),
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
    )


# ---------------------------------------------------------------------------
# 知识关联图
# ---------------------------------------------------------------------------


@router.get("/sessions/{session_id}/knowledge-links")
def get_knowledge_links(session_id: int, user=Depends(get_current_user)):
    """获取该 Session 的知识关联列表。"""
    from database import MindmapKnowledgeLink
    with db_session() as db:
        _assert_session_owner(db, session_id, user["id"])
        rows = (
            db.query(MindmapKnowledgeLink)
            .filter_by(user_id=user["id"], session_id=session_id)
            .all()
        )
        return [
            {
                "id": r.id,
                "source_node_id": r.source_node_id,
                "target_node_id": r.target_node_id,
                "source_node_text": r.source_node_text,
                "target_node_text": r.target_node_text,
                "link_type": r.link_type,
                "rationale": r.rationale,
            }
            for r in rows
        ]


@router.post("/sessions/{session_id}/knowledge-links/generate")
def generate_knowledge_links(session_id: int, user=Depends(get_current_user)):
    """用 LLM 从思维导图 Markdown 提取知识关联，覆盖旧数据。"""
    import json as _json
    from services.llm_service import LLMService
    from database import MindmapKnowledgeLink

    with db_session() as db:
        _assert_session_owner(db, session_id, user["id"])
        content = _get_mindmap_content(db, session_id)

    if not content or not content.strip():
        raise HTTPException(422, "思维导图内容为空，无法生成关联图")

    # 构建节点文本列表（去除性质前缀）
    node_texts = []
    for line in content.splitlines():
        m = re.match(r"^#{1,4}\s+(.*)", line)
        if m:
            text = m.group(1).strip()
            # 去除 ⭐⚠️🎯📌 前缀
            text = re.sub(r"^[⭐⚠️🎯📌]\s*", "", text).strip()
            if text:
                node_texts.append(text)

    if len(node_texts) < 3:
        raise HTTPException(422, "节点数量不足，无法生成关联图")

    prompt = (
        "你是一个知识图谱分析专家。请分析以下思维导图的节点列表，"
        "找出跨节点的知识关联关系。\n\n"
        f"节点列表：\n{chr(10).join(f'- {t}' for t in node_texts)}\n\n"
        "请识别 5-15 条最重要的跨节点关联，每条关联包含：\n"
        "- source_node_text：源节点文本（必须是上面列表中的节点）\n"
        "- target_node_text：目标节点文本（必须是上面列表中的节点）\n"
        "- link_type：关联类型，只能是 causal（因果）/ dependency（依赖）/ contrast（对比）/ evolution（演进）之一\n"
        "- rationale：关联依据，不超过 30 字\n\n"
        "以 JSON 数组格式输出，不要任何额外文字：\n"
        '[{"source_node_text":"...","target_node_text":"...","link_type":"...","rationale":"..."}]'
    )

    try:
        raw = LLMService().chat(
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
        )
    except RuntimeError as e:
        raise HTTPException(502, f"AI 服务暂时不可用：{e}")

    # 解析 JSON
    raw = raw.strip()
    if raw.startswith("```"):
        lines = raw.splitlines()
        raw = "\n".join(lines[1:-1] if lines[-1].strip() == "```" else lines[1:])
    try:
        items = _json.loads(raw)
        if not isinstance(items, list):
            raise ValueError("not a list")
    except (ValueError, _json.JSONDecodeError) as e:
        raise HTTPException(500, f"AI 返回格式异常：{e}")

    # 构建 node_id 映射（文本 → node_id）
    node_id_map: dict[str, str] = {}
    for line in content.splitlines():
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            raw_text = m.group(2).strip()
            clean_text = re.sub(r"^[⭐⚠️🎯📌]\s*", "", raw_text).strip()
            if clean_text:
                node_id_map[clean_text] = raw_text  # 用原始文本作为 node_id 近似

    # 写入数据库（覆盖旧数据）
    with db_session() as db:
        db.query(MindmapKnowledgeLink).filter_by(
            user_id=user["id"], session_id=session_id
        ).delete()

        valid_count = 0
        valid_types = {"causal", "dependency", "contrast", "evolution"}
        for item in items:
            src_text = str(item.get("source_node_text", "")).strip()
            dst_text = str(item.get("target_node_text", "")).strip()
            link_type = str(item.get("link_type", "")).strip()
            rationale = str(item.get("rationale", "")).strip()[:100]

            if not src_text or not dst_text or link_type not in valid_types:
                continue
            if src_text == dst_text:
                continue

            db.add(MindmapKnowledgeLink(
                user_id=user["id"],
                session_id=session_id,
                source_node_id=node_id_map.get(src_text, src_text),
                target_node_id=node_id_map.get(dst_text, dst_text),
                source_node_text=src_text,
                target_node_text=dst_text,
                link_type=link_type,
                rationale=rationale,
            ))
            valid_count += 1

    return {"ok": True, "count": valid_count}
