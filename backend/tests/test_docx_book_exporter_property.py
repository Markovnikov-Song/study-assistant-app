"""Property-based tests for DocxBookExporter.

# Feature: lecture-book-export, Property 8: Word ä¸­æå­ä½ä¸æ é¢æ ·å¼?
# Feature: lecture-book-export, Property 1: èç¹é¡ºåºä¿ç (docx)
"""
from __future__ import annotations

import sys
import os
import io

# Ensure backend package is importable when running from repo root
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", ".."))

from unittest.mock import patch, MagicMock

import pytest
from hypothesis import given, settings
from hypothesis import strategies as st

from docx import Document
from docx.oxml.ns import qn

from book_services.book_exporter import NodeInfo, TocEntry
from book_services.docx_book_exporter import DocxBookExporter


# ---------------------------------------------------------------------------
# Strategies
# ---------------------------------------------------------------------------

_node_text = st.text(
    alphabet=st.characters(
        whitelist_categories=("Lu", "Ll", "Nd"),
        whitelist_characters=" _-",
    ),
    min_size=1,
    max_size=30,
)

# XML-safe text: printable ASCII + common CJK range, no control chars
_xml_safe_text = st.text(
    alphabet=st.characters(
        whitelist_categories=("Lu", "Ll", "Nd", "Zs"),
        whitelist_characters=" _-.,!?",
        blacklist_categories=("Cc", "Cs"),  # no control chars, no surrogates
    ),
    min_size=1,
    max_size=30,
)

_heading_block = st.fixed_dictionaries({
    "type": st.just("heading"),
    "level": st.integers(min_value=1, max_value=3),
    "content": _xml_safe_text,
})

_non_heading_block = st.fixed_dictionaries({
    "type": st.sampled_from(["paragraph", "code", "list", "quote"]),
    "content": _xml_safe_text,
})

_any_block = st.one_of(_heading_block, _non_heading_block)


@st.composite
def nodes_with_all_block_types(draw, min_nodes=1, max_nodes=6):
    """Generate NodeInfo list where blocks cover all block types."""
    n = draw(st.integers(min_value=min_nodes, max_value=max_nodes))
    nodes = []
    for i in range(n):
        node_id = f"node_{i}"
        text = draw(_node_text)
        depth = draw(st.integers(min_value=1, max_value=4))
        # Ensure at least one block; mix heading and non-heading
        blocks = draw(st.lists(_any_block, min_size=1, max_size=6))
        nodes.append(NodeInfo(node_id=node_id, text=text, depth=depth, blocks=blocks))
    return nodes


@st.composite
def nodes_with_blocks(draw, min_nodes=1, max_nodes=8):
    """Generate NodeInfo list where every node has at least one block."""
    n = draw(st.integers(min_value=min_nodes, max_value=max_nodes))
    nodes = []
    for i in range(n):
        node_id = f"node_{i}"
        text = draw(_node_text)
        depth = draw(st.integers(min_value=1, max_value=4))
        blocks = draw(st.lists(_non_heading_block, min_size=1, max_size=4))
        nodes.append(NodeInfo(node_id=node_id, text=text, depth=depth, blocks=blocks))
    return nodes


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_ALLOWED_CHINESE_FONTS = {"å®ä½", "å¾®è½¯éé»"}
_ALLOWED_FONTS_LOWER = {f.lower() for f in _ALLOWED_CHINESE_FONTS}
# Code blocks use Courier New â?that is acceptable for non-Chinese content
_CODE_FONT = "courier new"


def _get_run_east_asian_font(run) -> str | None:
    """Extract the w:eastAsia font name from a run's rPr, if present."""
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return None
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        return None
    return rFonts.get(qn("w:eastAsia"))


def _is_acceptable_font(font_name: str | None, para_style: str) -> bool:
    """Return True if font_name is acceptable for the given paragraph style."""
    if font_name is None:
        return True  # no explicit font set â?inherits from style, acceptable
    fn_lower = font_name.lower()
    if fn_lower in _ALLOWED_FONTS_LOWER:
        return True
    # Code blocks legitimately use Courier New
    if fn_lower == _CODE_FONT:
        return True
    return False


def _build_docx_with_mock_latex(nodes, include_toc=False):
    """Call DocxBookExporter.build() with LatexRenderer mocked to return None."""
    import book_services.docx_book_exporter as _mod

    mock_renderer = MagicMock()
    mock_renderer.render.return_value = None  # triggers fallback text path

    with patch.object(_mod, "LatexRenderer", return_value=mock_renderer):
        exporter = DocxBookExporter()
        result = exporter.build(
            session_title="Test",
            nodes=nodes,
            include_toc=include_toc,
        )
    return result


# ---------------------------------------------------------------------------
# Property 8: Word ä¸­æå­ä½ä¸æ é¢æ ·å¼?
# Validates: Requirements 5.2, 5.3, 5.4
# ---------------------------------------------------------------------------

@given(nodes_with_all_block_types())
@settings(max_examples=50, deadline=None)
def test_word_chinese_fonts_and_heading_styles(nodes):
    """Property 8: Word ä¸­æå­ä½ä¸æ é¢æ ·å¼?

    For any block list containing heading/paragraph/code/list/quote blocks,
    DocxBookExporter.build() must produce a docx where:
    - All paragraph runs use å®ä½, å¾®è½¯éé», or Courier New (for code).
    - Heading blocks map to the correct Heading N paragraph style.

    Validates: Requirements 5.2, 5.3, 5.4
    # Feature: lecture-book-export, Property 8: Word ä¸­æå­ä½ä¸æ é¢æ ·å¼?
    """
    docx_bytes = _build_docx_with_mock_latex(nodes, include_toc=False)

    assert isinstance(docx_bytes, bytes) and len(docx_bytes) > 0

    doc = Document(io.BytesIO(docx_bytes))

    # Build a flat list of (block_type, level, content) for heading blocks
    # so we can correlate them with paragraphs in the document.
    heading_blocks: list[tuple[int, str]] = []  # (level, content)
    for node in nodes:
        for block in node.blocks:
            if block.get("type") == "heading":
                level = max(1, min(3, block.get("level", 1)))
                content = block.get("content", "") or block.get("text", "") or ""
                heading_blocks.append((level, content))

    # --- Font check: every run in every paragraph ---
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        for run in para.runs:
            # Check run.font.name (ASCII font)
            ascii_font = run.font.name
            if ascii_font is not None:
                assert _is_acceptable_font(ascii_font, style_name), (
                    f"Paragraph style={style_name!r}: run ASCII font {ascii_font!r} "
                    f"is not in allowed set {_ALLOWED_CHINESE_FONTS | {'Courier New'}}"
                )
            # Check East-Asian font via XML
            ea_font = _get_run_east_asian_font(run)
            if ea_font is not None:
                assert _is_acceptable_font(ea_font, style_name), (
                    f"Paragraph style={style_name!r}: run East-Asian font {ea_font!r} "
                    f"is not in allowed set {_ALLOWED_CHINESE_FONTS | {'Courier New'}}"
                )

    # --- Heading style check: heading blocks â?Heading N style ---
    # Collect paragraphs that have a Heading style
    heading_paras = [
        p for p in doc.paragraphs
        if p.style and p.style.name.startswith("Heading ")
    ]

    # Every heading block must correspond to a Heading N paragraph
    assert len(heading_paras) >= len(heading_blocks), (
        f"Expected at least {len(heading_blocks)} Heading paragraphs, "
        f"found {len(heading_paras)}"
    )

    # Verify each heading paragraph's style matches the block level
    heading_para_iter = iter(heading_paras)
    for level, content in heading_blocks:
        expected_style = f"Heading {level}"
        para = next(heading_para_iter)
        assert para.style.name == expected_style, (
            f"Heading block level={level} content={content!r}: "
            f"expected style {expected_style!r}, got {para.style.name!r}"
        )


# ---------------------------------------------------------------------------
# Property 1 (Docx): èç¹é¡ºåºä¿ç
# Validates: Requirements 3.1
# ---------------------------------------------------------------------------

@given(nodes_with_blocks())
@settings(max_examples=50, deadline=None)
def test_toc_order_preserves_input_order_docx(nodes):
    """Property 1 (Docx): èç¹é¡ºåºä¿ç

    For any ordered list of NodeInfo (all with blocks), _build_toc() must
    return TocEntry objects in the exact same order as the input nodes.

    Validates: Requirements 3.1
    # Feature: lecture-book-export, Property 1: èç¹é¡ºåºä¿ç (docx)
    """
    exporter = DocxBookExporter()
    toc_entries = exporter._build_toc(nodes)

    assert len(toc_entries) == len(nodes), (
        f"TOC has {len(toc_entries)} entries but input has {len(nodes)} nodes"
    )

    for i, (entry, node) in enumerate(zip(toc_entries, nodes)):
        assert entry.title == node.text, (
            f"Position {i}: TOC title {entry.title!r} != node text {node.text!r}. "
            f"Order is not preserved."
        )
        assert entry.depth == node.depth, (
            f"Position {i}: TOC depth {entry.depth} != node depth {node.depth}"
        )
        expected_anchor = f"node_{node.depth}_{node.text}"
        assert entry.anchor == expected_anchor, (
            f"Position {i}: anchor {entry.anchor!r} != expected {expected_anchor!r}"
        )
        assert entry.page == 0, (
            f"Position {i}: Word TOC page should be 0, got {entry.page}"
        )
